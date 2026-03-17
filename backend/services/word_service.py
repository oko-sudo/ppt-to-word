"""
Word(.docx) 문서 생성 서비스

- 2열 테이블 (슬라이드 번호 | 텍스트)
- 검수 항목: 빨간색 Bold
- 사용자 입력: 원문[입력값] 빨간색 Bold
- 삭제 항목: 텍스트에서 제거
"""
import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import PATTERNS
from storage.session_store import SessionData


# Word 출력 설정
WORD_CONFIG = PATTERNS["word_output"]
REVIEW_COLOR = RGBColor(
    int(WORD_CONFIG["review_item_color"][0:2], 16),
    int(WORD_CONFIG["review_item_color"][2:4], 16),
    int(WORD_CONFIG["review_item_color"][4:6], 16),
)
REVIEW_BOLD = WORD_CONFIG["review_item_bold"]
COL1_WIDTH = Cm(WORD_CONFIG["table_col1_width_cm"])
COL2_WIDTH = Cm(WORD_CONFIG["table_col2_width_cm"])

# 한글 음절 범위
KOREAN_RE = re.compile(
    r"[\uAC00-\uD7A3\u1100-\u11FF\u3130-\u318F]"
)


def generate_word(
    session: SessionData,
    decisions: list[dict],
) -> bytes:
    """
    Word 문서를 생성하고 bytes로 반환합니다.

    Args:
        session: 세션 데이터 (ppt_filename, prefix, slides 포함)
        decisions: 사용자 검수 결과 목록

    Returns:
        .docx 파일 bytes
    """
    doc = Document()

    # ── 제목 ────────────────────────────────────────────────────────────────
    title_text = Path(session.ppt_filename).stem  # 확장자 제외 파일명
    title_para = doc.add_heading(title_text, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 2열 테이블 생성 ──────────────────────────────────────────────────────
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # 헤더 행
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "슬라이드 번호"
    hdr_cells[1].text = "텍스트"
    _set_cell_bold(hdr_cells[0])
    _set_cell_bold(hdr_cells[1])

    # 검수 결정 맵 구축: {원문: {"user_input": ..., "delete": ...}}
    decision_map = {
        d["text"]: {"user_input": d.get("user_input", ""), "delete": d.get("delete", False)}
        for d in decisions
    }

    # ── 슬라이드 행 추가 ────────────────────────────────────────────────────
    slide_number_fmt = PATTERNS["word_output"]["slide_number_format"]

    for slide in session.slides:
        text = slide.cleaned_text
        if not text:
            continue

        slide_label = slide_number_fmt.format(
            prefix=session.prefix,
            number=slide.slide_number,
        )

        row_cells = table.add_row().cells

        # 열1: 슬라이드 번호
        row_cells[0].text = slide_label
        _center_cell(row_cells[0])

        # 열2: 텍스트 (비한글 항목에 서식 적용)
        _fill_text_cell(row_cells[1], text, decision_map)

    # 열 너비 설정
    _set_column_widths(table)

    # BytesIO로 저장 후 bytes 반환
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ── 셀 텍스트 채우기 (서식 적용) ────────────────────────────────────────────

def _fill_text_cell(cell, text: str, decision_map: dict) -> None:
    """
    테이블 셀에 텍스트를 채우고, 검수 항목에 서식을 적용합니다.

    알고리즘:
    1. 텍스트를 단락 단위로 분리
    2. 각 단락을 '검수 항목 / 일반 텍스트' 세그먼트로 분리
    3. 검수 항목 세그먼트에 빨간색 Bold 적용
    """
    # 기본 단락 제거
    cell._tc.clear_content()

    paragraphs = text.split("\n")

    for para_text in paragraphs:
        if not para_text.strip():
            continue

        para = cell.add_paragraph()
        segments = _split_by_review_items(para_text, decision_map)

        for seg_text, is_review, decision in segments:
            if not seg_text:
                continue

            if is_review:
                if decision.get("delete"):
                    continue  # 삭제 항목 건너뜀

                run = para.add_run()
                if decision.get("user_input"):
                    run.text = f"{seg_text}[{decision['user_input']}]"
                else:
                    run.text = seg_text

                # 빨간색 Bold 적용
                run.bold = REVIEW_BOLD
                run.font.color.rgb = REVIEW_COLOR
            else:
                para.add_run(seg_text)


def _split_by_review_items(
    text: str,
    decision_map: dict,
) -> list[tuple[str, bool, dict | None]]:
    """
    텍스트를 (세그먼트 텍스트, 검수항목 여부, decision) 튜플 목록으로 분리합니다.

    검수 항목이 여러 개일 경우 긴 것부터 먼저 매칭합니다 (부분 매칭 방지).
    """
    if not decision_map:
        return [(text, False, None)]

    # 긴 항목부터 먼저 매칭
    sorted_items = sorted(decision_map.keys(), key=len, reverse=True)
    pattern = "|".join(re.escape(item) for item in sorted_items)
    regex = re.compile(f"({pattern})")

    parts = regex.split(text)
    result = []

    for part in parts:
        if not part:
            continue
        if part in decision_map:
            result.append((part, True, decision_map[part]))
        else:
            result.append((part, False, None))

    return result


# ── 보조 함수 ───────────────────────────────────────────────────────────────

def _set_cell_bold(cell) -> None:
    """셀 텍스트를 Bold로 설정합니다."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
        # 텍스트가 run 없이 직접 설정된 경우
        if not para.runs and para.text:
            run = para.add_run(para.text)
            run.bold = True
            # 기존 직접 텍스트 제거
            for elem in para._p.findall(qn("w:r")):
                pass  # runs로 이미 처리됨


def _center_cell(cell) -> None:
    """셀 텍스트를 가운데 정렬합니다."""
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _set_column_widths(table) -> None:
    """테이블 열 너비를 설정합니다."""
    for row in table.rows:
        row.cells[0].width = COL1_WIDTH
        row.cells[1].width = COL2_WIDTH
